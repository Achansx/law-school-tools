"""Document parser for reading .docx files."""

from datetime import datetime
from pathlib import Path

from docx import Document

from ..models import DocumentMetadata, LegalDocument
from ..utils import get_logger

logger = get_logger(__name__)


class DocumentParseError(Exception):
    """Exception raised when document parsing fails."""

    pass


class DocumentParser:
    """Parser for Word (.docx) documents."""

    def __init__(self):
        """Initialize the document parser."""
        pass

    def parse(self, file_path: Path) -> LegalDocument:
        """Parse a .docx file into a LegalDocument object.

        Args:
            file_path: Path to the .docx file

        Returns:
            LegalDocument with extracted paragraphs

        Raises:
            DocumentParseError: If parsing fails
        """
        logger.info(f"Parsing document: {file_path}")

        try:
            doc = Document(str(file_path))
        except Exception as e:
            raise DocumentParseError(f"Failed to open document: {e}")

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append(text)

        if not paragraphs:
            raise DocumentParseError("Document contains no readable text")

        logger.info(f"Extracted {len(paragraphs)} paragraphs")

        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            total_paragraphs=len(paragraphs),
            processing_timestamp=datetime.now().isoformat(),
        )

        # Create document object
        legal_doc = LegalDocument(metadata=metadata, paragraphs=paragraphs)

        # Calculate word count
        metadata.total_words = legal_doc.count_words()
        logger.info(f"Total words: {metadata.total_words}")

        return legal_doc

    def extract_headings(self, file_path: Path) -> list[tuple[str, str]]:
        """Extract headings and their styles from the document.

        This can be useful for more advanced claim detection.

        Args:
            file_path: Path to the .docx file

        Returns:
            List of (heading_text, style_name) tuples
        """
        try:
            doc = Document(str(file_path))
        except Exception as e:
            logger.error(f"Failed to extract headings: {e}")
            return []

        headings = []
        for para in doc.paragraphs:
            style_name = para.style.name
            if "heading" in style_name.lower():
                text = para.text.strip()
                if text:
                    headings.append((text, style_name))

        logger.debug(f"Extracted {len(headings)} headings")
        return headings

    def get_document_info(self, file_path: Path) -> dict:
        """Get basic information about the document.

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary with document information
        """
        try:
            doc = Document(str(file_path))

            # Get core properties
            core_props = doc.core_properties

            info = {
                "title": core_props.title or "Untitled",
                "author": core_props.author or "Unknown",
                "created": core_props.created,
                "modified": core_props.modified,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            return info

        except Exception as e:
            logger.error(f"Failed to get document info: {e}")
            return {}
